import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import json
import os
import zipfile
import threading
import datetime
from decimal import Decimal
import traceback
import sys
import re  # 新增：用于解析SQL获取表名

# ---------- 依赖检查 ----------
try:
    import psycopg2
    import psycopg2.extras
except ImportError:
    messagebox.showerror("缺少依赖", "请安装 psycopg2：pip install psycopg2-binary")
    raise

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    messagebox.showerror("缺少依赖", "请安装 python-docx：pip install python-docx")
    raise

# ==================== 硬编码数据库配置 ====================
DB_CONFIG = {
    "host": "10.3.17.71",
    "port": 58324,
    "user": "damp",
    "password": "XGcmm3Bk5CTE3ehk7pKH",
    "database": "damp_prod",
    "schema": "data_service",
    "table": "DS_API"
}

DATA_APP_CONFIG = {
    "database": "damp_prod",
    "schema": "data_application"
}

DISPLAY_COLUMNS = ["API_URL", "API_VERSION", "NAME", "DEL_FLAG"]

# ==================== 数据库管理类 ====================
class DatabaseManager:
    def __init__(self, db_config):
        self.db_config = db_config
        self.conn = None
        self.conn_params = {
            "host": DB_CONFIG["host"],
            "port": DB_CONFIG["port"],
            "user": DB_CONFIG["user"],
            "password": DB_CONFIG["password"],
            "database": db_config["database"],
            "options": f"-c search_path={db_config.get('schema', 'public')}"
        }

    def connect(self):
        self.conn = psycopg2.connect(**self.conn_params)
        self.conn.set_client_encoding('UTF8')

    def close(self):
        if self.conn:
            self.conn.close()

    def get_columns(self, table_name):
        with self.conn.cursor() as cur:
            cur.execute("""
                SELECT column_name
                FROM information_schema.columns
                WHERE table_schema = %s AND table_name = %s
                ORDER BY ordinal_position
            """, (self.db_config.get("schema", "public"), table_name))
            return [row[0] for row in cur.fetchall()]

    def fetch_all(self, table_name):
        columns = self.get_columns(table_name)
        if not columns:
            raise ValueError(f"表 {table_name} 不存在或没有列")
        with self.conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute(f'SELECT * FROM "{self.db_config["schema"]}"."{table_name}" ORDER BY "NAME" DESC')
            rows = cur.fetchall()
        data = []
        for row in rows:
            data.append([row[col] for col in columns])
        return columns, data

class DataAppFetcher:
    """从 damp_prod.data_application 查询数据，支持自动重连"""
    def __init__(self):
        self.conn = None
        self.schema = DATA_APP_CONFIG["schema"]
        self.conn_params = {
            "host": DB_CONFIG["host"],
            "port": DB_CONFIG["port"],
            "user": DB_CONFIG["user"],
            "password": DB_CONFIG["password"],
            "database": DATA_APP_CONFIG["database"],
        }

    def connect(self):
        self.conn = psycopg2.connect(**self.conn_params)
        self.conn.set_client_encoding('UTF8')

    def close(self):
        if self.conn:
            self.conn.close()

    def fetch_one_row(self, table_name, fields):
        if not fields or not table_name:
            return []
        # 字段名可能包含大写，保险起见保留双引号
        cols = ', '.join(f'"{f}"' for f in fields)
        # 表名不强制双引号，以匹配实际存储的大小写
        query = f'SELECT {cols} FROM {self.schema}.{table_name} LIMIT 1'

        def _execute():
            with self.conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                cur.execute(query)
                row = cur.fetchone()
                return dict(row) if row else []

        try:
            if not self.conn or self.conn.closed:
                self.connect()
            return _execute()
        except Exception:
            try:
                self.connect()
                return _execute()
            except Exception:
                return []

# ==================== 主界面 ====================
class ApiDocGenerator:
    FIELD_MAP = {
        "name":          ["NAME", "API_NAME"],
        "version":       ["API_VERSION"],
        "url":           ["API_URL"],
        "method":        ["REQ_METHOD"],
        "config_json":   ["CONFIG_JSON"],
        "req_params":    ["REQ_PARAMS"],
        "res_params":    ["RES_PARAMS"],
        "description":   ["DESCRIPTION"],
        "del_flag":      ["DEL_FLAG"]
    }

    def __init__(self, root):
        self.root = root
        self.root.title(f"接口文档生成器 - {DB_CONFIG['schema']}.{DB_CONFIG['table']}")
        self.root.geometry("1000x700")

        self.db = DatabaseManager(DB_CONFIG)
        self.data_fetcher = None
        self.headers = []
        self.all_rows = []
        self.filtered_rows = []
        self.selected_iids = set()
        self.col_index = {}
        self.display_indices = []
        self.display_headers = []

        # 排序状态
        self.sort_column = None
        self.sort_ascending = True

        self._build_ui()
        self.status_var = tk.StringVar(value="就绪")
        ttk.Label(root, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W).pack(side=tk.BOTTOM, fill=tk.X)

    def _build_ui(self):
        # 数据库信息
        info_frame = ttk.Frame(self.root, padding=5)
        info_frame.pack(fill=tk.X, padx=10, pady=5)
        ttk.Label(info_frame, text=f"数据库: {DB_CONFIG['host']}:{DB_CONFIG['port']}/{DB_CONFIG['database']}",
                  font=('', 9, 'bold')).pack(side=tk.LEFT)
        ttk.Label(info_frame, text=f"模式: {DB_CONFIG['schema']} | 表: {DB_CONFIG['table']}",
                  font=('', 9)).pack(side=tk.LEFT, padx=20)

        # 加载按钮 + 映射状态
        btn_frame = ttk.Frame(self.root, padding=5)
        btn_frame.pack(fill=tk.X, padx=10, pady=5)
        ttk.Button(btn_frame, text="加载数据", command=self.load_data_thread).pack(side=tk.LEFT, padx=5)
        self.map_status = tk.StringVar(value="尚未加载")
        ttk.Label(btn_frame, textvariable=self.map_status, foreground="gray").pack(side=tk.LEFT, padx=10)

        # 筛选区域
        filter_frame = ttk.LabelFrame(self.root, text="数据筛选", padding=5)
        filter_frame.pack(fill=tk.X, padx=10, pady=5)

        ttk.Label(filter_frame, text="名称模糊:").pack(side=tk.LEFT, padx=5)
        self.filter_name = tk.StringVar()
        ttk.Entry(filter_frame, textvariable=self.filter_name, width=20).pack(side=tk.LEFT, padx=5)

        ttk.Label(filter_frame, text="删除标识:").pack(side=tk.LEFT, padx=5)
        self.filter_del = tk.StringVar()
        ttk.Entry(filter_frame, textvariable=self.filter_del, width=6).pack(side=tk.LEFT, padx=5)

        ttk.Label(filter_frame, text="URL模糊:").pack(side=tk.LEFT, padx=5)
        self.filter_url = tk.StringVar()
        ttk.Entry(filter_frame, textvariable=self.filter_url, width=20).pack(side=tk.LEFT, padx=5)

        ttk.Button(filter_frame, text="筛选", command=self.apply_filter).pack(side=tk.LEFT, padx=10)
        ttk.Button(filter_frame, text="重置", command=self.reset_filter).pack(side=tk.LEFT, padx=5)

        # 表格（含滚动条）
        mid_frame = ttk.LabelFrame(self.root, text="接口列表（勾选需导出的行）", padding=5)
        mid_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        self.select_all_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(mid_frame, text="全选 / 取消全选", variable=self.select_all_var,
                        command=self.toggle_all).pack(anchor=tk.W, pady=2)

        tree_frame = ttk.Frame(mid_frame)
        tree_frame.pack(fill=tk.BOTH, expand=True)

        scroll_y = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL)
        scroll_y.pack(side=tk.RIGHT, fill=tk.Y)
        scroll_x = ttk.Scrollbar(mid_frame, orient=tk.HORIZONTAL)
        scroll_x.pack(side=tk.BOTTOM, fill=tk.X)

        self.tree = ttk.Treeview(tree_frame, selectmode='none', show='headings',
                                 yscrollcommand=scroll_y.set,
                                 xscrollcommand=scroll_x.set)
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scroll_y.config(command=self.tree.yview)
        scroll_x.config(command=self.tree.xview)

        self.tree.bind('<Button-1>', self.on_tree_header_click)

        # 导出按钮
        export_frame = ttk.Frame(self.root, padding=5)
        export_frame.pack(fill=tk.X, padx=10, pady=5)
        ttk.Button(export_frame, text="生成 Word 文档", command=self.start_generate_docs).pack(side=tk.RIGHT, padx=10)

    # ---------- 数据加载 ----------
    def load_data_thread(self):
        self.status_var.set("正在连接数据库...")
        self.root.config(cursor="watch")
        threading.Thread(target=self.load_data, daemon=True).start()

    def load_data(self):
        try:
            self.db.close()
            self.db.connect()
            self.headers, self.all_rows = self.db.fetch_all(DB_CONFIG["table"])
            self._build_column_mapping()
            self._prepare_display_columns()
            self.filtered_rows = self.all_rows.copy()
            self.sort_column = "NAME"
            self.sort_ascending = False
            self.root.after(0, self._after_load)
        except Exception as e:
            self.root.after(0, lambda err=e: messagebox.showerror("错误", str(err)))
            self.root.after(0, lambda: self.status_var.set("加载失败"))
        finally:
            self.root.after(0, lambda: self.root.config(cursor=""))

    def _after_load(self):
        self.refresh_tree()
        self.selected_iids.clear()
        self.select_all_var.set(False)
        self.status_var.set(f"已加载 {len(self.all_rows)} 行数据")

    def _build_column_mapping(self):
        self.col_index = {}
        if not self.headers:
            self.map_status.set("⚠ 未加载列名")
            return
        status_parts = []
        for field, candidates in self.FIELD_MAP.items():
            for idx, col_name in enumerate(self.headers):
                if col_name.upper() in [c.upper() for c in candidates]:
                    self.col_index[field] = idx
                    status_parts.append(f"✅ {field}")
                    break
            else:
                status_parts.append(f"❌ {field}")
        self.map_status.set(" | ".join(status_parts))

    def _prepare_display_columns(self):
        self.display_indices = []
        self.display_headers = []
        for col_name in DISPLAY_COLUMNS:
            for idx, h in enumerate(self.headers):
                if h.upper() == col_name.upper():
                    self.display_indices.append(idx)
                    self.display_headers.append(h)
                    break

    def refresh_tree(self):
        self._sort_filtered_rows()

        self.tree.delete(*self.tree.get_children())
        if not self.filtered_rows or not self.display_headers:
            return

        col_ids = ["select"] + [f"col_{i}" for i in range(len(self.display_headers))]
        self.tree['columns'] = col_ids
        self.tree.heading("select", text="☐")
        self.tree.column("select", width=40, anchor=tk.CENTER, stretch=False)

        for i, header in enumerate(self.display_headers):
            symbol = " ↕"
            if header == self.sort_column:
                symbol = " ↑" if self.sort_ascending else " ↓"
            self.tree.heading(f"col_{i}", text=header + symbol)
            self.tree.column(f"col_{i}", width=150, anchor=tk.W)

        for i, row in enumerate(self.filtered_rows):
            values = ["☐"] + [str(row[idx]) if idx < len(row) and row[idx] is not None else "" for idx in self.display_indices]
            iid = self.tree.insert('', tk.END, values=values, iid=str(i))
            self.tree.item(iid, tags=('unselected',))
        self.tree.tag_configure('selected', background='#cce5ff')
        self.tree.tag_configure('unselected', background='white')

    def _sort_filtered_rows(self):
        if not self.sort_column or not self.filtered_rows:
            return
        try:
            col_idx = self.headers.index(self.sort_column)
        except ValueError:
            return

        def safe_val(row):
            val = row[col_idx] if col_idx < len(row) else ""
            if val is None:
                return ""
            return str(val).lower()

        self.filtered_rows.sort(key=safe_val, reverse=not self.sort_ascending)

    # ---------- 表头点击排序（三态循环） ----------
    def on_tree_header_click(self, event):
        region = self.tree.identify_region(event.x, event.y)
        if region != 'heading':
            if region == 'cell':
                self.on_tree_click(event)
            return

        column_id = self.tree.identify_column(event.x)
        if column_id == '#0':
            return
        col_num = int(column_id.replace('#', ''))
        if col_num == 1:
            return
        col_idx = col_num - 2
        if col_idx < 0 or col_idx >= len(self.display_headers):
            return

        clicked_header = self.display_headers[col_idx]

        if self.sort_column == clicked_header:
            if self.sort_ascending:
                self.sort_ascending = False
            else:
                self.sort_column = None
                self.sort_ascending = True
        else:
            self.sort_column = clicked_header
            self.sort_ascending = True

        self.refresh_tree()

    # ---------- 筛选 ----------
    def apply_filter(self):
        nf = self.filter_name.get().strip().lower()
        df = self.filter_del.get().strip()
        uf = self.filter_url.get().strip().lower()
        filtered = []
        for row in self.all_rows:
            if nf and nf not in self._get_cell(row, "name").lower(): continue
            if df and self._get_cell(row, "del_flag").strip() != df: continue
            if uf and uf not in self._get_cell(row, "url").lower(): continue
            filtered.append(row)
        self.filtered_rows = filtered
        self.refresh_tree()
        self.selected_iids.clear()
        self.select_all_var.set(False)
        self.status_var.set(f"筛选后显示 {len(self.filtered_rows)} 行")

    def reset_filter(self):
        self.filter_name.set("")
        self.filter_del.set("")
        self.filter_url.set("")
        self.filtered_rows = self.all_rows.copy()
        self.refresh_tree()
        self.selected_iids.clear()
        self.select_all_var.set(False)
        self.status_var.set(f"显示全部 {len(self.all_rows)} 行")

    # ---------- 行选择 ----------
    def on_tree_click(self, event):
        region = self.tree.identify_region(event.x, event.y)
        if region != 'cell': return
        column = self.tree.identify_column(event.x)
        if column != '#1': return
        iid = self.tree.identify_row(event.y)
        if not iid: return
        self.toggle_row(iid)

    def toggle_row(self, iid):
        tags = self.tree.item(iid, 'tags')
        if 'selected' in tags:
            self.tree.item(iid, tags=('unselected',))
            self.tree.set(iid, column='select', value='☐')
            self.selected_iids.discard(iid)
        else:
            self.tree.item(iid, tags=('selected',))
            self.tree.set(iid, column='select', value='☑')
            self.selected_iids.add(iid)
        total = len(self.tree.get_children())
        self.select_all_var.set(len(self.selected_iids) == total and total > 0)

    def toggle_all(self):
        is_all = self.select_all_var.get()
        for iid in self.tree.get_children():
            if is_all:
                self.tree.item(iid, tags=('selected',))
                self.tree.set(iid, column='select', value='☑')
                self.selected_iids.add(iid)
            else:
                self.tree.item(iid, tags=('unselected',))
                self.tree.set(iid, column='select', value='☐')
                self.selected_iids.discard(iid)

    def _get_cell(self, row, field):
        idx = self.col_index.get(field)
        if idx is not None and idx < len(row):
            val = row[idx]
            if val is None:
                return ""
            if isinstance(val, (list, dict)):
                return json.dumps(val, ensure_ascii=False)
            return str(val)
        return ""

    # ---------- 文档生成入口 ----------
    def start_generate_docs(self):
        if not self.selected_iids:
            messagebox.showwarning("提示", "请至少选择一行数据")
            return

        selected_rows = [self.filtered_rows[int(iid)] for iid in self.tree.get_children() if iid in self.selected_iids]
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

        if len(selected_rows) == 1:
            api_name = self._get_api_name(selected_rows[0])
            default_name = f"{api_name}_{timestamp}.docx"
            file_path = filedialog.asksaveasfilename(
                title="保存接口文档",
                defaultextension=".docx",
                initialfile=default_name,
                filetypes=[("Word 文档", "*.docx")]
            )
            if not file_path: return
            self.status_var.set("正在生成文档...")
            self.root.config(cursor="watch")
            threading.Thread(target=self._save_single_docx, args=(selected_rows[0], file_path), daemon=True).start()
        else:
            default_zip = f"接口文档_{timestamp}.zip"
            zip_path = filedialog.asksaveasfilename(
                title="保存接口文档压缩包",
                defaultextension=".zip",
                initialfile=default_zip,
                filetypes=[("ZIP 压缩文件", "*.zip")]
            )
            if not zip_path: return
            self.status_var.set("正在生成文档...")
            self.root.config(cursor="watch")
            threading.Thread(target=self._generate_zip_thread, args=(selected_rows, zip_path), daemon=True).start()

    def _save_single_docx(self, row, file_path):
        try:
            self._save_docx(row, file_path)
            self.root.after(0, lambda: messagebox.showinfo("成功", f"文档已保存到:\n{file_path}"))
            self.root.after(0, lambda: self.status_var.set("文档已保存"))
        except Exception as e:
            self.root.after(0, lambda err=e: messagebox.showerror("错误", str(err)))
        finally:
            self.root.after(0, lambda: self.root.config(cursor=""))

    def _generate_zip_thread(self, selected_rows, zip_path):
        try:
            local_fetcher = DataAppFetcher()
            local_fetcher.connect()
            self.data_fetcher = local_fetcher
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
                for row in selected_rows:
                    api_name = self._get_api_name(row)
                    tmp_path = f"temp_{api_name}.docx"
                    self._save_docx(row, tmp_path)
                    zf.write(tmp_path, f"{api_name}.docx")
                    os.remove(tmp_path)
            self.root.after(0, lambda: messagebox.showinfo("成功", f"文档已打包保存到:\n{zip_path}"))
            self.root.after(0, lambda: self.status_var.set(f"已生成 {len(selected_rows)} 个文档并打包"))
        except Exception as e:
            self.root.after(0, lambda err=e: messagebox.showerror("错误", str(err)))
        finally:
            if self.data_fetcher:
                self.data_fetcher.close()
                self.data_fetcher = None
            self.root.after(0, lambda: self.root.config(cursor=""))

    def _get_api_name(self, row):
        name = self._get_cell(row, "name")
        if not name:
            url = self._get_cell(row, "url")
            name = url.strip("/").split("/")[-1] if url else "api"
        return "".join(c for c in name if c not in r'\/:*?"<>|').strip() or "unnamed"

    def _build_url(self, row):
        version = self._get_cell(row, "version").strip("/")
        api_url = self._get_cell(row, "url").strip("/")
        base = "dataservicefront.chinacoalsx.com/services"
        if version: base += f"/{version}"
        if api_url: base += f"/{api_url}"
        return base

    def _convert_method(self, method_val):
        if method_val == "1": return "GET"
        elif method_val == "2": return "POST"
        return method_val

    # ---------- 安全参数提取 ----------
    def _safe_list(self, val):
        if isinstance(val, list):
            return [item for item in val if isinstance(item, dict)]
        return []

    def _get_req_params(self, row):
        idx = self.col_index.get("req_params")
        if idx is not None and idx < len(row):
            raw = row[idx]
            if isinstance(raw, list): return self._safe_list(raw)
            if isinstance(raw, str): return self._safe_list(self._parse_json(raw))
        cfg = self._get_config_dict(row)
        if cfg:
            req = cfg.get("reqParams")
            if req is not None: return self._safe_list(req)
        return []

    def _get_res_params(self, row):
        idx = self.col_index.get("res_params")
        if idx is not None and idx < len(row):
            raw = row[idx]
            if isinstance(raw, list): return self._safe_list(raw)
            if isinstance(raw, str): return self._safe_list(self._parse_json(raw))
        cfg = self._get_config_dict(row)
        if cfg:
            res = cfg.get("resParams")
            if res is not None: return self._safe_list(res)
        return []

    def _get_config_dict(self, row):
        idx = self.col_index.get("config_json")
        if idx is not None and idx < len(row):
            raw = row[idx]
            if isinstance(raw, dict): return raw
            if isinstance(raw, str):
                try: return json.loads(raw)
                except: return {}
        return {}

    def _get_config_value(self, row, key):
        cfg = self._get_config_dict(row)
        return cfg.get(key) if cfg else None

    # ---------- 从 SQL 中提取表名 ----------
    def _extract_table_from_sql(self, sql_text):
        """从 CONFIG_JSON 的 sqlText 字段中提取 FROM 后的表名"""
        if not sql_text:
            return None
        # 1. 去除注释、换行等，将多行SQL合并为单行（保留空格）
        sql = ' '.join(sql_text.split())
        # 2. 匹配 FROM 后跟着的表名（可能包含 schema.table 或直接表名）
        #    忽略 ${where} 等动态条件，只取第一个遇到的表名
        match = re.search(r'\bfrom\s+([a-zA-Z0-9_."]+)', sql, re.IGNORECASE)
        if match:
            table = match.group(1).strip()
            # 去掉可能的双引号
            table = table.replace('"', '')
            return table
        return None

    def _get_table_name(self, row):
        """从 CONFIG_JSON 中提取表名（优先 tableName，其次 sqlText）"""
        # 直接尝试 tableName（虽然你说没有，但保留以兼容未来）
        cfg = self._get_config_dict(row)
        if cfg:
            # 若有 tableName 键则使用
            table = cfg.get("tableName")
            if table:
                return table
            # 否则从 sqlText 中提取
            sql = cfg.get("lastSqlText")
            if sql:
                return self._extract_table_from_sql(sql)
        return None

    def _parse_json(self, text):
        if not text: return []
        try:
            obj = json.loads(text)
            return obj if isinstance(obj, list) else []
        except: return []

    # ---------- 生成 docx ----------
    def _save_docx(self, row, file_path):
        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(10.5)

        try:
            name = self._get_cell(row, "name") or "接口"
            version = self._get_cell(row, "version")
            full_url = self._build_url(row)
            method = self._convert_method(self._get_cell(row, "method"))
            desc = self._get_cell(row, "description")

            document.add_heading(name, level=1)
            document.add_paragraph(f"版本: {version}")
            document.add_paragraph(f"URL: {full_url}")
            document.add_paragraph(f"请求方法: {method}")
            if desc:
                document.add_paragraph(f"描述: {desc}")

            # 固定参数
            document.add_heading('固定参数', level=2)
            fixed_params = [
                {"paramName": "pageSize", "paramComment": "-", "nullable": "1", "defaultValue": "20", "exampleValue": "-"},
                {"paramName": "pageNum", "paramComment": "-", "nullable": "1", "defaultValue": "1", "exampleValue": "-"}
            ]
            ftable = document.add_table(rows=1, cols=5, style='Light Grid Accent 1')
            fhdr = ftable.rows[0].cells
            fhdr[0].text = '字段名'; fhdr[1].text = '描述'; fhdr[2].text = '允许为空'; fhdr[3].text = '默认值'; fhdr[4].text = '示例值'
            for p in fixed_params:
                row_cells = ftable.add_row().cells
                row_cells[0].text = p["paramName"]
                row_cells[1].text = p["paramComment"]
                row_cells[2].text = "是" if p["nullable"] == "1" else "否"
                row_cells[3].text = p["defaultValue"]
                row_cells[4].text = p["exampleValue"]

            # 请求参数
            req_params = self._get_req_params(row)
            document.add_heading('请求参数', level=2)
            if req_params:
                table = document.add_table(rows=1, cols=5, style='Light Grid Accent 1')
                hdr = table.rows[0].cells
                hdr[0].text = '字段名'; hdr[1].text = '描述'; hdr[2].text = '允许为空'; hdr[3].text = '默认值'; hdr[4].text = '示例值'
                for p in req_params:
                    if not isinstance(p, dict): continue
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(p.get("paramName", ""))
                    row_cells[1].text = str(p.get("paramComment", ""))
                    row_cells[2].text = "是" if str(p.get("nullable")) == "1" else "否"
                    row_cells[3].text = str(p.get("defaultValue", ""))
                    row_cells[4].text = str(p.get("exampleValue", ""))
            else:
                document.add_paragraph("无")

            # 响应参数
            res_params = self._get_res_params(row)
            document.add_heading('响应参数', level=2)
            if res_params:
                table = document.add_table(rows=1, cols=4, style='Light Grid Accent 1')
                hdr = table.rows[0].cells
                hdr[0].text = '字段名'; hdr[1].text = '说明'; hdr[2].text = '类型'; hdr[3].text = '必填'
                for p in res_params:
                    if not isinstance(p, dict): continue
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(p.get("fieldName", p.get("engName", "")))
                    row_cells[1].text = str(p.get("fieldComment", p.get("cnName", "")))
                    row_cells[2].text = str(p.get("columnType", ""))
                    row_cells[3].text = "是" if str(p.get("nullableFlag")) == "1" else "否"
            else:
                document.add_paragraph("无")

            # 响应数据示例
            data = self._get_response_data(row)
            response_json = self._format_response_json(data)
            document.add_heading('响应数据', level=2)
            p = document.add_paragraph()
            run = p.add_run(response_json)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

        except Exception as e:
            document.add_heading('文档生成出错', level=1)
            document.add_paragraph(f"错误信息: {str(e)}\n\n{traceback.format_exc()}")
            print(f"[ERROR] {traceback.format_exc()}", file=sys.stderr)

        document.save(file_path)

    def _get_response_data(self, row):
        # 使用新的表名提取逻辑
        table_name = self._get_table_name(row)
        if not table_name:
            return []
        res_params = self._get_res_params(row)
        fields = [p.get("fieldName") for p in res_params if isinstance(p, dict) and p.get("fieldName")]
        if not fields:
            return []
        if not self.data_fetcher:
            self.data_fetcher = DataAppFetcher()
            self.data_fetcher.connect()
        row_data = self.data_fetcher.fetch_one_row(table_name, fields)
        return [row_data] if row_data else []

    def _format_response_json(self, data):
        class DecimalEncoder(json.JSONEncoder):
            def default(self, obj):
                if isinstance(obj, Decimal):
                    return float(obj)
                return super().default(obj)

        template = {
            "msg": "操作成功",
            "code": 200,
            "data": {
                "pageNum": 1,
                "pageSize": 20,
                "total": len(data),
                "data": data
            }
        }
        return json.dumps(template, ensure_ascii=False, indent=2, cls=DecimalEncoder)

if __name__ == "__main__":
    root = tk.Tk()
    app = ApiDocGenerator(root)
    root.mainloop()